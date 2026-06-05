from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = r"E:\work\MoyaMixMatrix\business_docs\tongjian-systematic-business-plan-v1.1.docx"

FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.2):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_fixed(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    total_dxa = sum(int(w * 1440) for w in widths)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_style_font(style, size, color=None, bold=None):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(r_fonts)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_para(doc, text="", size=10.5, color=None, bold=False, italic=False, align=None, before=0, after=6, line=1.2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, before, after, line)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_spacing(p, after=3, line=1.18)
    r = p.add_run(text)
    set_run_font(r, size=10.3)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_spacing(p, after=4, line=1.18)
    r = p.add_run(text)
    set_run_font(r, size=10.3)
    return p


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table, [6.5])
    set_table_borders(table, color="D7DEE8", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4, line=1.18)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.18)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.2, color=RGBColor(35, 35, 35))
    add_para(doc, "", after=2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_fixed(table, widths)
    set_table_borders(table)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")
    hdr = table.rows[0].cells
    for idx, label in enumerate(headers):
        cell = hdr[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.12)
            if idx == 0 and len(headers) <= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=font_size, color=RGBColor(35, 35, 35), bold=(idx == 0))
    add_para(doc, "", after=2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        set_paragraph_spacing(p, before=14, after=8, line=1.15)
    elif level == 2:
        set_paragraph_spacing(p, before=10, after=5, line=1.15)
    else:
        set_paragraph_spacing(p, before=8, after=4, line=1.15)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 10.5, RGBColor(25, 25, 25))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        set_style_font(style, size, color, bold=True)
        style.paragraph_format.keep_with_next = True
    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        set_style_font(style, 10.3, RGBColor(35, 35, 35))
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18
    return doc


def add_cover(doc):
    add_para(doc, "商业计划书", size=12, color=GRAY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "彤见矩阵", size=30, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.05)
    add_para(
        doc,
        "AI 短视频增长操盘系统",
        size=16,
        color=DARK_BLUE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )
    add_para(
        doc,
        "OPC 转化 · 企业培训 · 认证代理 · 授权渠道 · 系统消耗平台",
        size=11.5,
        color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=28,
    )
    add_callout(
        doc,
        "核心口号",
        "让每个卖点都被看见。彤见矩阵的机会不是再做一个剪辑工具，而是把短视频经营从“凭感觉做内容”变成“围绕利益点持续测试、复盘和复制”的系统。",
        fill="EEF5FF",
    )
    metadata = [
        ("创始人", "杨彤"),
        ("版本", "系统商业叙事 v1.1"),
        ("日期", "2026 年 5 月 31 日"),
        ("文档用途", "融资沟通、渠道沟通、团队内部战略对齐"),
    ]
    add_table(doc, ["项目", "内容"], metadata, [1.5, 5.0], header_fill=LIGHT_BLUE, font_size=9.8)
    add_para(
        doc,
        "重要说明：本计划书中的公开市场数据均注明来源；涉及营收、漏斗、成本、转化率的数字为阶段目标或测算假设，正式融资材料需要用试点实测数据替换。",
        size=9.5,
        color=GRAY,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=16,
        after=0,
    )
    doc.add_page_break()


def build_doc():
    doc = setup_document()
    add_cover(doc)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("彤见矩阵 | 系统性商业计划书 v1.1")
    set_run_font(r, size=9, color=GRAY)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = ""
    r = footer.add_run("让每个卖点都被看见")
    set_run_font(r, size=9, color=GRAY)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. 执行摘要", 1)
    add_para(
        doc,
        "彤见矩阵不是单纯靠终端客户订阅赚钱的 AI 剪辑软件，而是一套围绕商业短视频增长能力搭建的“产品 + 方法 + 渠道”系统。它先用真实用户的内容操盘场景验证产品，再通过 OPC 转化、企业培训、认证代理、授权渠道、系统服务费和 AI 消耗形成分阶段收入。",
    )
    add_para(
        doc,
        "产品当前方向是 AI 短视频增长操盘系统：帮助有销售、获客、曝光和转化需求的人，完成从卖点提炼、内容生成、批量测试、发布台账到数据复盘的闭环。适用对象不局限于实体门店，也包括电商商家、品牌方、矩阵号团队、MCN/代运营公司、达人主播和不会剪辑但需要内容的人。",
    )
    add_callout(
        doc,
        "一句话商业判断",
        "彤见矩阵早期用终端客户和 OPC 跑出案例与现金流，中期用企业培训和认证代理把方法规模化，后期用授权渠道、系统服务费、AI 消耗和数据资产放大平台价值。",
    )
    add_bullet(doc, "对投资人：这不是“又一个 AI 剪辑工具”，而是把培训、渠道和系统消耗串成收入链的 AI 短视频增长平台。")
    add_bullet(doc, "对渠道方：它不是只卖软件，而是给服务商一套可教学、可交付、可复购、可质检的业务包。")
    add_bullet(doc, "对终端用户：它解决的不是会不会剪辑，而是今天卖什么、怎么说利益点、发完怎么判断有效。")

    add_heading(doc, "2. 数据调研与市场依据", 1)
    add_para(
        doc,
        "公开数据可以证明三个大前提：短视频是基础流量入口，线上消费和本地生活的经营数字化仍在扩大，AI 内容合规正在成为商家使用门槛。真正需要试点证明的是：哪类用户最愿意持续付费、哪个内容场景使用频率最高、代理和企业培训是否能把收入放大。",
    )
    market_rows = [
        (
            "短视频用户",
            "截至 2025 年 12 月，短视频用户规模达 10.74 亿，占网民整体 95.4%。",
            "短视频已经是商业内容的基础入口，商家需要的是持续生产和测试能力。",
            "CNNIC 第57次报告",
        ),
        (
            "网络购物用户",
            "截至 2025 年 12 月，网络购物用户规模达 9.37 亿，占网民整体 83.2%。",
            "电商、品牌、新品推广和投放素材测试具备长期内容需求。",
            "CNNIC 第57次报告",
        ),
        (
            "网上零售",
            "2025 年全国网上零售额 15.97 万亿元，同比增长 8.6%。",
            "商品视频、详情页转短视频、卖点 A/B 测试有可服务空间。",
            "国家统计局/CNNIC",
        ),
        (
            "餐饮消费",
            "2025 年餐饮收入 57982 亿元，同比增长 3.2%。",
            "本地生活到店、团购套餐和门店活动适合做高频短视频内容测试。",
            "国家统计局",
        ),
        (
            "餐饮连锁化",
            "2025 年餐厅总数约 747 万家，餐饮连锁化率达 25%。",
            "连锁和区域品牌更适合标准化内容 SOP、代理交付和系统复用。",
            "CCFA/美团白皮书",
        ),
        (
            "AI 合规",
            "《人工智能生成合成内容标识办法》自 2025 年 9 月 1 日起施行。",
            "AI 视频产品必须处理标识、授权、审核和平台规则，合规能力可成为差异点。",
            "国家网信办",
        ),
    ]
    add_table(doc, ["数据点", "公开数据", "对项目的意义", "来源"], market_rows, [1.1, 2.1, 2.3, 1.0], header_fill=LIGHT_BLUE, font_size=8.4)
    add_heading(doc, "2.1 市场机会的真实切口", 2)
    add_para(
        doc,
        "项目不能只用“短视频市场大”作为论据。彤见矩阵真正要切的是商业内容生产中的高频、重复、可模板化、可训练、可复盘环节：商品卖点测试、门店活动推广、矩阵号批量发布、客户内容交付、直播切片复用和无剪辑基础用户的内容启动。",
    )
    add_para(
        doc,
        "因此，商业计划书的验证重点应从宏观市场规模转向一手经营数据：OPC 线索成本、训练营转化率、企业培训客单价、代理续费率、AI 生成消耗、终端内容发布率和复盘后复购率。",
    )

    add_heading(doc, "3. 目标人群与真实痛点", 1)
    add_para(
        doc,
        "彤见矩阵的产品适用人群不应被限制为实体商家。更准确的定义是：所有需要通过短视频完成销售、引流、获客、曝光或内容交付的人。早期可以选择一个高频场景做验证，但品牌叙事要站在商业短视频增长系统上。",
    )
    user_rows = [
        (
            "电商商家 / 品牌方",
            "商品图转视频、详情页转短视频、商品口播、新品推广、卖点测试、多版本投放。",
            "上新快、素材多、投放需要多版本，但内部缺少持续脚本和剪辑产能。",
            "系统订阅、AI 美豆、企业培训、投放素材包。",
        ),
        (
            "本地生活商家",
            "门店图转视频、活动海报转视频、团购套餐推广、同城引流、节日促销。",
            "知道要发短视频，但不会持续选题、提炼利益点和复盘团购效果。",
            "7天试点包、内容测试包、培训营、代理交付。",
        ),
        (
            "短视频运营 / 矩阵号团队",
            "批量混剪、极速裂变、多账号发布、A/B 测试、标题封面优化、成片库管理。",
            "账号多、需求快、内容变量多，人工剪辑和管理成本高。",
            "团队版系统、AI 消耗、成片库、发布台账。",
        ),
        (
            "MCN / 代运营 / 内容营销公司",
            "客户素材管理、行业模板复用、爆款结构参考、统一包装、批量出片、数据复盘。",
            "客户多、交付重、利润被人工和返修吃掉，难以复制稳定 SOP。",
            "认证代理、系统服务费、模板授权、质检督导。",
        ),
        (
            "达人 / 主播 / 个人 IP",
            "直播切片、数字人口播、AI 字幕、课程推广、商品讲解、IP 内容复用。",
            "内容资产分散，直播和长内容无法高效二次分发。",
            "训练营、AI 美豆、数字人、切片工具包。",
        ),
        (
            "不会剪辑但需要内容的人",
            "上传素材、选择场景和模板，即可生成可发布视频。",
            "缺少剪辑技能，也不知道商业短视频该讲什么利益点。",
            "低价公开课、实战营、模板会员、基础订阅。",
        ),
    ]
    add_table(doc, ["人群", "高频用途", "真实痛点", "收入入口"], user_rows, [1.25, 2.0, 2.1, 1.15], header_fill=LIGHT_BLUE, font_size=8.0)
    add_callout(
        doc,
        "焦点原则",
        "品牌目标人群可以宽，但 60 天验证切口必须窄。建议优先选择“本地生活团购内容测试”或“电商商品卖点素材测试”其中一个跑出案例，再复制到其他人群。",
    )

    add_heading(doc, "4. 当前产品方向", 1)
    add_para(
        doc,
        "当前产品应聚焦为“AI 短视频增长操盘系统”，不是泛剪辑工具。用户进入系统后，不是先面对时间轴和复杂编辑器，而是先回答商业问题：这次要卖什么、卖点是什么、给谁看、用什么结构表达、发布后看什么数据。",
    )
    add_heading(doc, "4.1 OPC 产品逻辑", 2)
    add_para(doc, "产品底层采用 OPC 操盘逻辑，把短视频生产拆成三个可训练、可系统化的环节：")
    add_number(doc, "Offer：明确本轮要推广的商品、套餐、活动、课程、直播间或服务。")
    add_number(doc, "Point：提炼用户能看懂、能被打动的利益点，包括价格、效果、场景、信任、稀缺、便利、身份感。")
    add_number(doc, "Content：围绕不同利益点生成多条短视频、标题、封面文案、口播稿和发布建议，并通过数据复盘继续迭代。")
    add_heading(doc, "4.2 MVP 功能边界", 2)
    mvp_rows = [
        ("素材输入", "商品图、门店图、海报、直播片段、详情页、课程资料、客户素材。"),
        ("卖点提炼", "按行业和场景自动拆解利益点、目标人群、转化理由和风险表达。"),
        ("内容测试包", "一次生成多个卖点方向、短视频版本、标题封面和 CTA，用于 A/B 测试。"),
        ("批量生成", "支持图文转视频、口播视频、数字人、混剪变体、直播切片和模板化包装。"),
        ("合规审核", "检查价格、活动有效期、夸大承诺、肖像授权、AI 标识、行业敏感词。"),
        ("发布台账", "记录账号、发布时间、素材版本、标题、卖点方向、链接和后续数据。"),
        ("复盘迭代", "回收曝光、点击、私信、转化线索、核销或人工填报数据，给出下一轮建议。"),
    ]
    add_table(doc, ["模块", "说明"], mvp_rows, [1.35, 5.15], header_fill=LIGHT_GRAY, font_size=9.2)
    add_heading(doc, "4.3 北极星指标", 2)
    add_callout(
        doc,
        "产品北极星指标",
        "每个付费用户每月完成“创建 Offer -> 提炼 Point -> 生成 Content -> 发布 -> 复盘 -> 再生成”的有效循环次数。",
        fill="EEF5FF",
    )
    for metric in [
        "生成后发布率：避免产品停留在“生成玩具”，验证内容是否真的进入经营场景。",
        "发布后数据回收率：验证系统能否从工具走向操盘，形成迭代依据。",
        "复盘后再次生成率：验证用户是否把系统当成持续内容工厂，而不是一次性体验。",
        "代理交付活跃率：验证渠道是否真正使用系统交付，而不是只买授权。",
        "AI 美豆复购率：验证高频生成是否成立，避免消耗模式只是转嫁成本。",
    ]:
        add_bullet(doc, metric)

    add_heading(doc, "5. 最终方向与竞争优势", 1)
    add_para(
        doc,
        "最终方向不是单点 AI 视频工具，而是 AI 短视频增长服务行业的操作系统：底层是内容生成和数据复盘，中层是课程、模板、SOP 和质检，上层是企业培训、认证代理和授权渠道。这个方向的壁垒不来自某一个数字人或模板，而来自真实案例、行业利益点库、渠道交付标准和系统消耗网络的叠加。",
    )
    comp_rows = [
        (
            "剪映 / 模板工具",
            "便宜、易上手，但仍依赖人写脚本、找卖点、做复盘。",
            "彤见矩阵从 Offer 和利益点开始，不只是剪辑操作。",
        ),
        (
            "通用 AI 视频工具",
            "画面生成能力强，但不懂商品、套餐、转化路径和交付 SOP。",
            "彤见矩阵把商业短视频拆成可测试的卖点和内容包。",
        ),
        (
            "数据平台",
            "能看榜单和趋势，但不能直接生成下一轮内容。",
            "彤见矩阵把趋势、模板和生成流程连接到发布复盘。",
        ),
        (
            "代运营公司",
            "能交付内容，但人力重、质量不透明、难复制。",
            "彤见矩阵让服务商用统一系统、模板、台账和质检交付。",
        ),
        (
            "培训机构",
            "能教方法，但课后工具和交付链路弱。",
            "彤见矩阵把培训结果导入系统使用、代理认证和持续消耗。",
        ),
    ]
    add_table(doc, ["对比对象", "客户现有选择的问题", "彤见矩阵的差异"], comp_rows, [1.45, 2.55, 2.5], header_fill=LIGHT_BLUE, font_size=8.4)
    add_heading(doc, "5.1 可量化优势指标", 2)
    add_para(
        doc,
        "现阶段不能直接声称“效果领先”，应把优势写成可验证指标。融资前建议用 30-50 个试点账号或商家补齐下列数据。",
    )
    proof_rows = [
        ("内容效率", "从素材上传到 20 条可发布视频的平均用时。", "目标：低于传统人工交付 70% 以上。"),
        ("发布率", "生成内容中真实发布的比例。", "目标：超过 50%，证明不是低价值生成。"),
        ("复盘率", "发布后能回收至少一项数据的内容比例。", "目标：超过 40%，逐步提升自动化。"),
        ("返修率", "客户要求重做或人工大改的比例。", "目标：随模板和审核规则下降。"),
        ("代理人效", "一个代理团队每月可服务客户数和内容量。", "目标：比纯人工代运营提升 2-3 倍。"),
        ("美豆复购", "付费用户主动充值生成额度的比例。", "目标：验证高频刚需和消耗收入。"),
    ]
    add_table(doc, ["优势方向", "衡量方式", "验证目标"], proof_rows, [1.3, 2.7, 2.5], header_fill=LIGHT_GRAY, font_size=8.5)

    add_heading(doc, "6. 商业模式总述", 1)
    add_para(
        doc,
        "彤见矩阵不是只挣终端客户的钱。它的商业结构应该是多环节收益，但每一环必须服务同一条主线：用内容操盘方法获客，用培训建立信任，用终端案例证明有效，用代理承接重交付，用授权渠道放大规模，用系统服务费和 AI 消耗形成持续收入。",
    )
    add_callout(
        doc,
        "商业主线",
        "OPC 内容获客 -> 低价训练营筛选需求 -> 企业培训/终端试点建立案例 -> 认证代理承接重交付 -> 授权渠道复制 -> 系统服务费与 AI 美豆持续消耗 -> 企业版、白标、API、数据价值。",
    )
    revenue_rows = [
        ("前端转化", "OPC 公开课、直播课、低价训练营", "商家、创业者、服务商、小白用户", "获客、筛选、建立认知"),
        ("高客单信任", "企业培训、内训、顾问项目", "品牌方、企业、商协会、机构", "现金流、背书、组织客户"),
        ("案例验证", "终端试点包、标杆操盘、内容测试包", "商家、品牌、达人团队", "产生真实案例和模板"),
        ("渠道建设", "认证代理营、服务商认证、城市/行业授权", "代运营、MCN、本地服务商、培训机构", "把重交付外移"),
        ("持续收入", "系统服务费、模板授权、AI 美豆、质检督导", "代理、企业、商家", "提升毛利和复购"),
        ("远期价值", "白标/API、数据报告、交易归因、分成", "企业客户、平台、渠道方", "平台化和数据化"),
    ]
    add_table(doc, ["收入层级", "收入项", "付费方", "存在目的"], revenue_rows, [1.1, 1.8, 1.8, 1.8], header_fill=LIGHT_BLUE, font_size=8.3)

    add_heading(doc, "7. 分阶段目标", 1)
    add_para(
        doc,
        "商业计划书需要明确主收入曲线。早期不把所有模式同时做重，而是用阶段推进：先验证真实需求和获客，再产品化方法，然后渠道化交付，最后系统化和平台化。",
    )

    stages = [
        (
            "阶段一：0-3 个月，OPC 获客与终端场景验证",
            "验证用户是否真的需要 AI 短视频操盘，并用低成本 OPC 获得商家、服务商和个人用户线索。",
            [
                "选择一个核心验证切口：本地生活团购内容测试或电商商品卖点素材测试。",
                "完成 20-50 个付费试点，形成 3-5 个可公开展示的案例。",
                "跑通公开课、训练营、试点包三个前端产品。",
            ],
            [
                "OPC 公开课：9.9-99 元。",
                "7-14 天实战营：399-999 元。",
                "终端内容测试包：999-2999 元。",
                "标杆操盘项目：1-3 万元。",
            ],
            [
                "OPC 线索 1000-3000 条。",
                "到课率 30%-50%。",
                "付费实战营 100-300 人。",
                "月收入目标 10-30 万元。",
            ],
        ),
        (
            "阶段二：3-6 个月，企业培训与方法产品化",
            "把一线操盘案例沉淀成课程、SOP、模板和企业训练体系，提升客单价和信任背书。",
            [
                "形成 OPC 方法课、行业案例课、卖点训练课、内容测试 SOP。",
                "交付 3-10 家企业培训或行业内训。",
                "把培训学员导入系统试用和后续代理认证。",
            ],
            [
                "企业内训：2-20 万元/场或项目。",
                "21 天企业实战营：1999-5999 元/人。",
                "行业模板包：999-9999 元。",
                "年度顾问：5-30 万元/年。",
            ],
            [
                "企业客户 3-10 家。",
                "可复用课程模块 10 个以上。",
                "行业案例 10-30 个。",
                "月收入目标 30-80 万元。",
            ],
        ),
        (
            "阶段三：6-12 个月，认证代理与重交付外移",
            "总部不再重度承接所有终端交付，而是输出系统、课程、模板、SOP、质检和品牌，由认证代理承接本地服务。",
            [
                "建立认证代理营和考试机制。",
                "上线代理工作台、客户台账、模板库、质检后台。",
                "明确禁止承诺清单、内容审核标准和分润规则。",
            ],
            [
                "认证代理营：1-5 万元/期。",
                "代理系统服务费：999-9999 元/月。",
                "AI 美豆：按量计费。",
                "模板授权和质检督导：按月或按项目。",
            ],
            [
                "认证代理 50-150 家。",
                "代理系统活跃率 70% 以上。",
                "代理贡献收入超过总收入 30%。",
                "平台月收入目标 80-200 万元。",
            ],
        ),
        (
            "阶段四：12-24 个月，授权渠道与系统消耗规模化",
            "把代理模型升级为城市、行业和企业授权体系，让收入从项目型转向授权费、系统费、AI 消耗和分润。",
            [
                "推出城市合伙人、行业合伙人、企业内训授权。",
                "建立总部质检、渠道督导和结算系统。",
                "用数据看板追踪代理交付质量和系统消耗。",
            ],
            [
                "城市合伙人：10-50 万元/年。",
                "行业合伙人：10-100 万元/年。",
                "企业内训授权：2-20 万元/项目。",
                "白标系统：10-100 万元/年。",
            ],
            [
                "授权城市/行业合伙人 5-20 个。",
                "认证代理 150-500 家。",
                "代理贡献收入超过总收入 50%。",
                "平台月收入目标 200-500 万元。",
            ],
        ),
        (
            "阶段五：24 个月以后，企业版、白标、API 与数据价值",
            "在渠道网络、系统使用和案例数据有规模后，升级为 AI 短视频商业内容基础设施。",
            [
                "推出企业版、白标版、API 接入和行业数据报告。",
                "探索交易归因、团购核销、线索分成和投放素材优化。",
                "将内容模板、行业卖点和效果数据形成持续更新的增长资产。",
            ],
            [
                "企业版系统订阅。",
                "白标/API 收入。",
                "行业数据会员。",
                "交易归因和线索分成。",
            ],
            [
                "企业客户续费率。",
                "API 调用量。",
                "数据报告复购率。",
                "系统消耗收入占比持续提升。",
            ],
        ),
    ]
    for title, goal, actions, revenue, kpis in stages:
        add_heading(doc, title, 2)
        add_para(doc, goal)
        add_table(
            doc,
            ["关键动作", "收入结构", "阶段指标"],
            [("\n".join(actions), "\n".join(revenue), "\n".join(kpis))],
            [2.15, 2.15, 2.2],
            header_fill=LIGHT_GRAY,
            font_size=8.4,
        )

    add_heading(doc, "8. 财务模型与融资用途", 1)
    add_para(
        doc,
        "财务模型建议按“训练营/培训现金流 + 代理授权现金流 + 系统消耗持续收入”拆开，而不是把所有收入混成一条 SaaS 订阅曲线。当前版本应明确：下列数字是测算目标，需要用试点数据验证。",
    )
    finance_rows = [
        ("启动期", "OPC、训练营、终端试点", "10-30 万元", "验证需求、沉淀案例、跑通漏斗"),
        ("验证期", "企业培训、标杆项目、代理认证", "30-80 万元", "提高客单价，筛选渠道和企业客户"),
        ("放大期", "授权渠道、系统费、AI 美豆、质检督导", "100-300 万元", "重交付外移，持续收入提升"),
        ("平台期", "企业版、白标、API、数据/分成", "300 万元以上", "平台化、数据化、行业基础设施"),
    ]
    add_table(doc, ["阶段", "收入结构", "月收入目标", "核心目的"], finance_rows, [1.1, 2.0, 1.25, 2.15], header_fill=LIGHT_BLUE, font_size=8.6)
    add_heading(doc, "8.1 单位经济模型需要补齐的数据", 2)
    for item in [
        "OPC 线索成本、到课率、完课率、训练营转化率。",
        "终端内容测试包的真实交付成本：脚本、素材整理、生成、人工审核、返修、客服、复盘。",
        "企业培训获客成本、交付人力、复购率和转介绍率。",
        "代理从认证到首单的周期、首月签约客户数、系统活跃率和投诉率。",
        "AI 单条生成成本、高清导出成本、云存储成本和美豆复购率。",
        "系统服务费、AI 消耗和授权费在总收入中的占比变化。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.2 融资建议", 2)
    add_para(doc, "建议天使轮融资目标：300-500 万元人民币。资金重点不应只投研发，也要投向 OPC 漏斗、培训产品化、代理渠道验证和合规体系。")
    funding_rows = [
        ("OPC 与内容获客", "20%", "短视频、直播、公开课、投放测试、线索转化。"),
        ("课程与企业培训体系", "15%", "课程开发、讲师体系、案例库、企业交付材料。"),
        ("代理渠道建设", "25%", "认证体系、招商、督导、区域试点、渠道物料。"),
        ("产品中台研发", "25%", "操盘台、代理工作台、生成系统、质检、结算和数据看板。"),
        ("合规与运营", "15%", "合同、授权、AI 标识、审核规则、客服和财务。"),
    ]
    add_table(doc, ["用途", "比例", "说明"], funding_rows, [1.6, 0.8, 4.1], header_fill=LIGHT_GRAY, font_size=8.8)

    add_heading(doc, "9. 风险与验证重点", 1)
    risk_rows = [
        ("目标人群过宽", "品牌叙事可以宽，但试点切口必须窄；60 天只验证一个高频场景。"),
        ("培训收入降低科技估值", "培训必须导向系统使用、代理认证和企业组织客户，而不是一次性卖课。"),
        ("内容交付变成人力外包", "总部只做标杆案例和 SOP，规模交付逐步交给认证代理和系统自动化。"),
        ("代理乱承诺影响品牌", "建立认证、合同、质检、督导、禁止承诺清单和投诉处理机制。"),
        ("数据回流困难", "早期手工复盘，中期表格导入和授权接入，后期 API/服务商系统打通。"),
        ("AI 内容合规风险", "做 AI 标识、肖像授权、价格有效期、夸大宣传、行业敏感表达审核。"),
        ("数字人和模板容易复制", "壁垒不能写成“有数字人/有模板”，而要写成真实投放验证后的行业利益点库和渠道使用数据。"),
    ]
    add_table(doc, ["风险", "应对方式"], risk_rows, [1.65, 4.85], header_fill=LIGHT_BLUE, font_size=8.8)
    add_heading(doc, "9.1 90 天必须拿到的证据", 2)
    for item in [
        "至少 20 个真实付费试点，记录行业、客单价、交付内容、发布数据、复购意向。",
        "至少 3 个可公开展示的案例，包含原本成本、生成效率、发布结果和客户反馈。",
        "一套 OPC 漏斗数据：曝光、留资、到课、付费训练营、试点包、代理咨询。",
        "一套代理验证数据：候选代理数量、付费认证率、首单周期、系统活跃率。",
        "一套生成成本数据：单条视频成本、人工审核时长、返修率、客服成本。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. 对外融资叙事", 1)
    add_para(
        doc,
        "彤见矩阵的融资叙事应从“我有一个 AI 剪辑软件”升级为“我在做 AI 短视频增长服务行业的产品化和渠道化基础设施”。",
    )
    add_callout(
        doc,
        "对外表达版本",
        "中小商家、品牌方和内容服务商未来不会只买一个剪辑工具，他们需要一套能持续发现利益点、生成内容、测试结果、复盘优化并被服务商规模化交付的 AI 内容工厂。彤见矩阵先用终端场景和 OPC 跑出真实案例，再通过企业培训、认证代理和授权渠道，把重交付转化为系统化收入和数据资产。",
        fill="EEF5FF",
    )
    add_para(
        doc,
        "这套叙事比单纯讲功能更适合融资：它承认早期需要服务和培训拿现金流，但把培训、代理和授权都纳入系统使用、AI 消耗和数据回流，而不是彼此割裂的赚钱方式。",
    )

    add_heading(doc, "11. 数据来源", 1)
    sources = [
        ("CNNIC 第57次《中国互联网络发展状况统计报告》", "https://www.cnnic.net.cn/NMediaFile/2026/0304/MAIN1772588317069TUXN3827X8.pdf"),
        ("国家统计局：2025年12月份社会消费品零售总额增长0.9%", "https://www.stats.gov.cn/sj/zxfbhjd/202601/t20260119_1962323.html"),
        ("CCFA：《2026中国餐饮连锁化发展白皮书》发布信息", "https://www.ccfa.org.cn/portal/cn/xiangxi.jsp?id=447139&type=10003"),
        ("国家网信办：《人工智能生成合成内容标识办法》答记者问", "https://www.cac.gov.cn/2025-03/14/c_1743654685896173.htm"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, line=1.15)
        r = p.add_run(label + "：")
        set_run_font(r, size=9.5, color=RGBColor(35, 35, 35), bold=True)
        add_hyperlink(p, url, url)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
